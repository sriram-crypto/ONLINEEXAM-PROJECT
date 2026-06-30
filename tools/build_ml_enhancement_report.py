from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\onlineexam\reports\ExamPulse_ML_and_DB_Enhancement_Report.docx")


COLORS = {
    "navy": "172033",
    "blue": "2754D8",
    "teal": "0F8F83",
    "green": "2F855A",
    "amber": "D98A16",
    "coral": "D84C5F",
    "soft": "F6F9FC",
    "line": "D9E2EC",
    "muted": "647087",
    "white": "FFFFFF",
}


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_callout(doc, title, body, fill="EAF7F5", accent="0F8F83"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cell_width(table.cell(0, 0), 0.15)
    set_cell_width(table.cell(0, 1), 6.05)
    shade_cell(table.cell(0, 0), accent)
    shade_cell(table.cell(0, 1), fill)
    table.cell(0, 0).text = ""
    cell = table.cell(0, 1)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    p.paragraph_format.space_after = Pt(3)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for i, header in enumerate(headers):
        shade_cell(header_row.cells[i], COLORS["navy"])
        set_cell_text(header_row.cells[i], header, bold=True, color=COLORS["white"], size=8.3)
        if widths:
            set_cell_width(header_row.cells[i], widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i < len(cells):
                shade_cell(cells[i], COLORS["white"] if len(table.rows) % 2 else COLORS["soft"])
                set_cell_text(cells[i], value, size=font_size)
                if widths:
                    set_cell_width(cells[i], widths[i])

    doc.add_paragraph()
    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("ExamPulse ML Enhancement Implementation Report")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.12
    for style_name, size, color in [
        ("Title", 24, COLORS["navy"]),
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12, COLORS["teal"]),
        ("Heading 3", 10.5, COLORS["navy"]),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    # Cover page
    cover = doc.add_table(rows=1, cols=1)
    cover.autofit = False
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_width(cover.cell(0, 0), 6.6)
    shade_cell(cover.cell(0, 0), COLORS["navy"])
    c = cover.cell(0, 0)
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ExamPulse")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(COLORS["white"])
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ML and Database Enhancement Implementation Report")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor.from_string("CFFAF4")
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Prepared for M.Tech Project Documentation")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor.from_string("E6F1F8")
    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Implementation date: 05 May 2026 | Workspace: C:\\onlineexam")
    r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor.from_string("D9E2EC")

    doc.add_paragraph()
    add_callout(
        doc,
        "Purpose of this document",
        "This report summarizes the ML enhancement work implemented in ExamPulse and explains the algorithm used for each enhancement in a format suitable for project review and viva discussion.",
        fill="F6F9FC",
        accent=COLORS["blue"],
    )

    meta_rows = [
        ("Backend ML service", "backend/services/mlEngine.js"),
        ("Backend ML APIs", "backend/routes/mlRoutes.js"),
        ("Frontend ML assistant", "frontend/src/components/MLQuestionAssistant/index.js"),
        ("Student ML experience", "frontend/src/layouts/student/myexam.js and MyResults.js"),
        ("Parent alerts", "frontend/src/layouts/parent/ParentPanel.js"),
        ("Superadmin ML risk", "frontend/src/layouts/superadmin/activateOrDeactivateExams.js"),
        ("Worksheet optimizer", "frontend/src/layouts/teacher/worksheets.js"),
        ("Full DB table usage", "backend/routes/superadmin/dbTables.js and frontend/src/layouts/superadmin/DatabaseTables.js"),
    ]
    add_table(doc, ["Area", "Main implementation file"], meta_rows, widths=[2.0, 4.4], font_size=8.5)

    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "The ExamPulse platform was enhanced with a lightweight, transparent ML layer that works directly with the existing MySQL database and React/Express application. "
        "The implementation avoids external model downloads and instead uses explainable algorithms: text similarity, weighted heuristics, trend analysis, entropy scoring, and greedy optimization. "
        "This makes the system suitable for an academic project because every prediction can be explained with visible formulas and data points."
    )
    bullet(doc, "Adaptive question recommendation was added to the student exam flow.")
    bullet(doc, "Question bank intelligence now supports auto-tagging, duplicate detection, quality scoring, and semantic search.")
    bullet(doc, "Student analytics now include predicted score, weak chapter heatmap, and recommended study focus.")
    bullet(doc, "Parent dashboard now shows predictive at-risk alerts and helps prepare meeting requests.")
    bullet(doc, "Superadmin exam report now shows ML proctoring/anomaly risk.")
    bullet(doc, "Teacher worksheet generation now includes AI Balanced selection.")
    bullet(doc, "A complete DB Tables module was added so the superadmin can use every live table from the perfect DB dump.")

    add_heading(doc, "2. Feature and Algorithm Matrix", 1)
    matrix_rows = [
        (
            "Text preprocessing foundation",
            "Normalization, stop-word removal, tokenization, term-frequency vectors.",
            "HTML removal, lowercase conversion, punctuation cleanup, word filtering, and token counting are used before similarity and classification.",
        ),
        (
            "Adaptive exam engine",
            "Weighted ability estimation plus difficulty bucket matching.",
            "Recent answer correctness and answer time estimate ability; next question is selected from Easy, Medium, or Hard pools based on target ability.",
        ),
        (
            "Smart performance analytics",
            "Weighted score prediction with linear trend slope.",
            "Uses recent average, latest score, regression slope, and attempt count to produce predicted score, confidence, and range.",
        ),
        (
            "Weak chapter detection",
            "Chapter-level accuracy aggregation and threshold classification.",
            "Counts correct vs attempted answers per chapter and classifies each topic as Strong, Practice More, or Needs Revision.",
        ),
        (
            "Question auto-tagging",
            "Rule-based difficulty inference and semantic metadata matching.",
            "Uses complexity cues, option count, text length, and text similarity against subjects/chapters to suggest level, type, subject, and chapter.",
        ),
        (
            "Duplicate question detection",
            "Hybrid cosine similarity and Jaccard similarity with exact-match boost.",
            "Flags a duplicate when semantic similarity is 0.85 or higher.",
        ),
        (
            "Semantic question search",
            "Hybrid semantic similarity with difficulty and type boosts.",
            "Searches question text plus metadata and ranks rows by semantic score.",
        ),
        (
            "Question quality scoring",
            "Weighted readability, distractor quality, and grammar signal.",
            "Final score = readability 45 percent + distractor quality 35 percent + grammar 20 percent.",
        ),
        (
            "Proctoring/anomaly risk",
            "Rule-based risk scoring with answer-pattern entropy.",
            "Combines tab switches, copy/paste events, very fast answers, score gap, and low entropy answer pattern.",
        ),
        (
            "Parent predictive alerts",
            "At-risk scoring from trend, latest score, and skipped ratio.",
            "Classifies whether a linked ward needs attention and recommends a meeting when risk crosses the threshold.",
        ),
        (
            "AI worksheet generator",
            "Greedy stratified selection with difficulty split and weak-chapter priority.",
            "Selects a balanced paper using Easy 40 percent, Medium 40 percent, Hard 20 percent by default, then fills gaps by priority.",
        ),
        (
            "Full database table usage",
            "Metadata-driven schema discovery and guarded generic CRUD.",
            "Reads information_schema, builds table/column metadata dynamically, supports safe search, sort, pagination, export, add, edit, and delete.",
        ),
    ]
    add_table(doc, ["Enhancement", "Algorithm Used", "How It Works"], matrix_rows, widths=[1.7, 2.0, 2.75], font_size=7.5)

    doc.add_page_break()

    add_heading(doc, "3. ML Algorithms Explained", 1)
    feature_details = [
        (
            "3.1 Text Similarity Engine",
            [
                "Algorithm: hybrid cosine similarity and Jaccard similarity.",
                "Used in: duplicate detection, semantic search, auto-tagging, chapter matching, and distractor quality.",
                "Formula idea: textSimilarity = max(exact match, cosine * 0.68 + Jaccard * 0.32).",
                "Reason: cosine similarity rewards repeated important terms, while Jaccard rewards shared concept words. Combining both gives stable ranking for short exam questions.",
            ],
        ),
        (
            "3.2 Difficulty Inference",
            [
                "Algorithm: heuristic complexity scoring.",
                "Signals: number of tokens, option count, complex cues such as derive, analyze, prove, compare, evaluate, and basic cues such as define, identify, true/false.",
                "Output: Easy, Medium, or Hard with confidence.",
                "Reason: question difficulty labels are often rule-driven in academic platforms, so transparent heuristics are easier to defend than a black-box classifier.",
            ],
        ),
        (
            "3.3 Adaptive Exam Selection",
            [
                "Algorithm: weighted ability estimation followed by difficulty bucket matching.",
                "Ability = 72 percent weighted accuracy + 28 percent recent accuracy. Later answers get slightly higher weight, and very quick correct answers get a small speed boost.",
                "Difficulty target: ability >= 0.72 means Hard, ability <= 0.42 means Easy, otherwise Medium.",
                "Next question is selected by smallest distance from target difficulty, with a freshness tie-breaker.",
            ],
        ),
        (
            "3.4 Score Prediction",
            [
                "Algorithm: recency-weighted prediction with linear regression slope.",
                "Uses last six attempts, average score, latest score, trend slope, and count of available results.",
                "Formula idea: predicted = average * 0.58 + latest * 0.28 + trend * 1.6 + attempt-count adjustment.",
                "Output includes predicted score, range, confidence, and explanation factors.",
            ],
        ),
        (
            "3.5 Question Quality Score",
            [
                "Algorithm: weighted quality scoring.",
                "Readability penalizes very long sentences and difficult long words.",
                "Distractor quality checks whether wrong options are plausible but not duplicates.",
                "Grammar signal checks capitalization, punctuation, repeated spaces, and repeated words.",
                "Final score is converted to grade A, B, C, or D.",
            ],
        ),
        (
            "3.6 Risk and Anomaly Detection",
            [
                "Algorithm: deterministic risk score plus entropy.",
                "Risk points come from tab switches, copy/paste events, very fast average answer time, score gap above baseline, and low answer-pattern entropy.",
                "Risk level: Low below 40, Medium from 40 to 69, High from 70 and above.",
                "Reason: exam proctoring must be explainable; each risk factor can be shown to administrators.",
            ],
        ),
        (
            "3.7 Worksheet Optimization",
            [
                "Algorithm: greedy stratified sampling.",
                "Default split: Easy 40 percent, Medium 40 percent, Hard 20 percent.",
                "Questions from weak chapters are prioritized first; remaining slots are filled by newest/available questions.",
                "Output includes selected question IDs, focus areas, and difficulty mix.",
            ],
        ),
    ]
    for heading, bullets in feature_details:
        if heading.startswith("3.7"):
            doc.add_page_break()
        add_heading(doc, heading, 2)
        for item in bullets:
            bullet(doc, item)

    add_heading(doc, "4. Backend Enhancements", 1)
    backend_rows = [
        ("backend/services/mlEngine.js", "Core ML utility module containing all explainable algorithms."),
        ("backend/routes/mlRoutes.js", "API endpoints for ML features such as tagging, search, prediction, adaptive selection, risk, and worksheet optimization."),
        ("backend/app.js", "Mounted /api/ml and /api/superadmin/db-tables routes."),
        ("backend/routes/student/myexam.js", "Stores timing metadata and attempt order for adaptive/risk analysis."),
        ("backend/routes/teacher/worksheets.js", "Returns chapter and difficulty metadata needed by worksheet optimizer."),
        ("backend/migrations/20260505_ml_enhancements.sql", "Adds optional ML quality, metadata, risk, timing, and prediction log fields."),
        ("backend/routes/superadmin/dbTables.js", "New metadata-driven APIs to expose every database table safely."),
    ]
    add_table(doc, ["Backend file", "Enhancement"], backend_rows, widths=[2.5, 3.95], font_size=8.2)

    add_heading(doc, "5. Frontend Enhancements", 1)
    frontend_rows = [
        ("MLQuestionAssistant", "Auto-tagging, duplicate warning, and quality grade while admin/teacher creates a question."),
        ("Admin QuestionTable", "Semantic Search button and AI-ranked question results."),
        ("Teacher QuestionTable", "Semantic Search button and AI-ranked question results."),
        ("Student MyExam", "Predicted score badges, adaptive recommendation during exam, timing data capture."),
        ("Student MyResults", "Expected next score, revision heatmap, and recommended study hours."),
        ("ParentPanel", "ML at-risk alerts for linked wards and pre-filled meeting request preparation."),
        ("Superadmin Activate Exams", "ML risk score and level shown in exam report and export."),
        ("Teacher Worksheets", "AI Balanced worksheet mode with focus areas and selected question IDs."),
        ("Superadmin DB Tables", "New screen to inspect, search, export, add, edit, and delete live table rows."),
    ]
    add_table(doc, ["Frontend area", "User-facing enhancement"], frontend_rows, widths=[2.05, 4.4], font_size=8.2)

    doc.add_page_break()
    add_heading(doc, "6. API Endpoint Summary", 1)
    endpoint_rows = [
        ("POST /api/ml/suggest-tags", "Auto-suggest level, type, subject, and chapter."),
        ("POST /api/ml/check-duplicate", "Detect similar questions using hybrid text similarity."),
        ("POST /api/ml/quality-score", "Return readability, distractor quality, grammar, total score, and grade."),
        ("POST /api/ml/search-questions", "Semantic question search across question text and metadata."),
        ("POST /api/ml/next-question", "Adaptive next-question recommendation."),
        ("GET /api/ml/predict-score", "Student score prediction and confidence range."),
        ("GET /api/ml/weak-chapters", "Chapter accuracy, weak topics, heatmap, and study-hour recommendation."),
        ("GET /api/ml/at-risk-students", "Parent alert risk classification."),
        ("GET /api/ml/risk-score", "Exam attempt anomaly/proctoring risk."),
        ("POST /api/ml/optimize-worksheet", "Balanced worksheet question selection."),
        ("GET /api/superadmin/db-tables", "Discover all live database tables."),
        ("GET/POST/PUT/DELETE /api/superadmin/db-tables/:table", "Browse and safely maintain rows for each table."),
    ]
    add_table(doc, ["Endpoint", "Purpose"], endpoint_rows, widths=[2.65, 3.8], font_size=8.0)

    doc.add_page_break()

    add_heading(doc, "7. Database and Full Table Usage Enhancement", 1)
    doc.add_paragraph(
        "The perfect database dump contains 55 live tables, and the live MySQL database was checked against it. The DB Tables module was added so superadmins can use all tables even if a separate business-specific screen does not exist yet."
    )
    db_rows = [
        ("Schema discovery", "Reads information_schema.TABLES and information_schema.COLUMNS dynamically."),
        ("Table grouping", "Uses keyword classification to group tables as Question Bank, Exam Engine, People, Access, Practice, Intelligence, Academic Setup, or General."),
        ("Row search", "Builds LIKE search across searchable text, enum, set, and JSON columns."),
        ("Safe sorting", "Sort column must exist in table metadata before it is used in SQL."),
        ("Safe mutation", "Add/edit/delete requires a known table, valid identifier, and primary key for update/delete."),
        ("Read-only protection", "Audit/log/event tables are protected from UI mutation."),
        ("Export", "Visible rows are exported as CSV from the frontend."),
    ]
    add_table(doc, ["DB enhancement", "Algorithm or technique"], db_rows, widths=[2.1, 4.35], font_size=8.2)

    add_heading(doc, "8. Verification Completed", 1)
    verification_rows = [
        ("Backend syntax", "node --check passed for ML engine, ML routes, student exam route, DB table route, and app.js."),
        ("Frontend build", "npm.cmd run build completed successfully."),
        ("Database match", "Dump and live DB both contained 55 tables with no missing or extra live tables."),
        ("ML smoke test", "Question quality and score prediction functions returned expected structured output."),
        ("Known warnings", "Existing React ESLint warnings remain in older files, but they do not block compilation."),
    ]
    add_table(doc, ["Check", "Result"], verification_rows, widths=[1.75, 4.7], font_size=8.2)

    add_heading(doc, "9. Viva Explanation Points", 1)
    numbered(doc, "The ML layer is explainable and deterministic, so every recommendation can be traced to visible features and weights.")
    numbered(doc, "The project uses hybrid text similarity instead of a black-box model for question search and duplicate detection because exam questions are short and structured.")
    numbered(doc, "Adaptive exams use ability estimation and difficulty buckets, which is a standard computer-based testing idea simplified for this project.")
    numbered(doc, "Risk scoring is not a final cheating decision; it is an alerting score to help administrators review suspicious attempts.")
    numbered(doc, "The DB Tables enhancement ensures complete database utilization by exposing all tables through metadata-driven APIs and a superadmin screen.")
    numbered(doc, "Future work can replace individual heuristic modules with trained models, such as TF-IDF vectorizers, logistic regression, random forest, or transformer embeddings, while preserving the same API structure.")

    add_callout(
        doc,
        "Final note",
        "This implementation provides practical ML-style intelligence without requiring external services or heavy model deployment. It is suitable for academic demonstration because the logic is transparent, testable, and connected end-to-end from database to backend API to frontend screen.",
        fill="FFF7E8",
        accent=COLORS["amber"],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_document()
    print(path)
